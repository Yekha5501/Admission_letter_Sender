import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import time
from datetime import datetime

# ---------- CONFIG ----------
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587
GMAIL_USER = "yekhapmandindi@gmail.com"
GMAIL_PASS = "cjpgbpainjcersvu"

# ---------- PATHS ----------
TEMPLATE_PATH = "admission_template.docx"
OUTPUT_DIR = "admission_letters"
LOG_FILE = "admission_log.xlsx"  # <-- NEW: Excel log file

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# ---------- NEW: FUNCTION TO WRITE LOG ----------
def write_to_log(student_data, pdf_generated, email_sent, error_message=""):
    """
    Appends a row to the Excel log file
    """
    # Prepare log data
    log_entry = {
        'Timestamp': datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        'Name': student_data['Name'],
        'Email': student_data['Email'],
        'Program': student_data.get('Program', ''),
        'PDF Generated': 'Yes' if pdf_generated else 'No',
        'Email Sent': 'Yes' if email_sent else 'No',
        'Status': 'Success' if (pdf_generated and email_sent) else 'Failed',
        'Error Message': error_message
    }
    
    # If log file exists, append to it
    if os.path.exists(LOG_FILE):
        # Read existing log
        df_log = pd.read_excel(LOG_FILE)
        # Append new entry
        df_log = pd.concat([df_log, pd.DataFrame([log_entry])], ignore_index=True)
    else:
        # Create new log
        df_log = pd.DataFrame([log_entry])
    
    # Save back to Excel
    df_log.to_excel(LOG_FILE, index=False)
    print(f"   📝 Logged to: {LOG_FILE}")

# ---------- FUNCTION: Generate PDF ----------
def generate_admission_letter(student_data):
    name = student_data['Name']
    
    doc = Document(TEMPLATE_PATH)
    
    replacements = {
        '{Name}': student_data['Name'],
        '{Program}': student_data['Program'],
        '{Faculty}': student_data['Faculty'],
        '{Duration}': student_data['Duration'],
        '{StartDate}': student_data['StartDate'],
        '{Date}': time.strftime("%B %d, %Y")
    }
    
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
    
    temp_docx = os.path.join(OUTPUT_DIR, f"{name}_admission.docx")
    doc.save(temp_docx)
    print(f"   📄 Word doc created: {name}_admission.docx")
    
    pdf_path = os.path.join(OUTPUT_DIR, f"{name}_admission.pdf")
    
    try:
        convert(temp_docx, pdf_path)
        print(f"   ✅ PDF generated: {name}_admission.pdf")
        return pdf_path, True, ""
    except Exception as e:
        error_msg = str(e)
        print(f"   ❌ PDF conversion failed: {error_msg}")
        return temp_docx, False, error_msg

# ---------- FUNCTION: Send Email ----------
def send_admission_email(student_data, pdf_path):
    name = student_data['Name']
    email = student_data['Email']
    
    msg = MIMEMultipart()
    msg['From'] = GMAIL_USER
    msg['To'] = email
    msg['Subject'] = f"Admission Letter - {name}"
    
    body = f"""
    Dear {name},
    
    Please find attached your admission letter from our institution.
    We are delighted to welcome you to our academic community.
    
    For any queries, please contact the Registry Office.
    
    Regards,
    Registry Office
    """
    msg.attach(MIMEText(body, 'plain'))
    
    try:
        with open(pdf_path, "rb") as attachment:
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f"attachment; filename={os.path.basename(pdf_path)}")
            msg.attach(part)
        print(f"   📎 Attached: {os.path.basename(pdf_path)}")
    except Exception as e:
        error_msg = str(e)
        print(f"   ❌ Attachment error: {error_msg}")
        return False, error_msg
    
    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(GMAIL_USER, GMAIL_PASS)
        server.send_message(msg)
        server.quit()
        print(f"   ✅ EMAIL SENT to {email}")
        return True, ""
    except Exception as e:
        error_msg = str(e)
        print(f"   ❌ Email failed: {error_msg}")
        return False, error_msg

# ---------- MAIN PROCESS ----------
def main():
    # Load student data
    df = pd.read_excel("students.xlsx")
    
    print(f"📊 Loaded {len(df)} students")
    print(f"📋 Columns: {df.columns.tolist()}")
    print("\n🚀 Starting admission letter generation...\n")
    
    success_count = 0
    fail_count = 0
    
    for index, row in df.iterrows():
        name = row['Name']
        print(f"📝 Processing: {name}")
        
        # Step 1: Generate PDF
        pdf_path, pdf_generated, pdf_error = generate_admission_letter(row)
        
        # Step 2: Send email
        if pdf_generated:
            email_sent, email_error = send_admission_email(row, pdf_path)
        else:
            email_sent = False
            email_error = "PDF generation failed, email not sent"
        
        # Step 3: Write to log
        error_message = ""
        if not pdf_generated:
            error_message = pdf_error
        if not email_sent and email_error:
            error_message = error_message + " | " + email_error if error_message else email_error
        
        write_to_log(row, pdf_generated, email_sent, error_message)
        
        # Track success/failure
        if pdf_generated and email_sent:
            success_count += 1
        else:
            fail_count += 1
        
        print("   " + "-"*40 + "\n")
        time.sleep(2)
    
    # Final summary
    print(f"✅ All done!")
    print(f"   ✅ Success: {success_count}")
    print(f"   ❌ Failed: {fail_count}")
    print(f"   📊 Log file saved to: {LOG_FILE}")

if __name__ == "__main__":
    main()