import pandas as pd
from docx import Document
from docx2pdf import convert
import os
import time
from datetime import datetime

# ---------- CONFIG ----------
TEMPLATE_PATH = "admission_template.docx"
OUTPUT_DIR = "admission_letters_pdf_only"

if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

# ---------- GENERATE PDFS ----------
def generate_pdf(student_data):
    name = student_data['Name']
    
    # Load the template
    doc = Document(TEMPLATE_PATH)
    
    # ONLY replace {Name} - nothing else
    replacements = {
        '{Name}': name
    }
    
    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)
    
    # Also check tables (if your template has them)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, value)
    
    # Save as Word doc (temporary)
    temp_docx = os.path.join(OUTPUT_DIR, f"{name}_admission.docx")
    doc.save(temp_docx)
    
    # Convert to PDF
    pdf_path = os.path.join(OUTPUT_DIR, f"{name}_admission.pdf")
    
    try:
        convert(temp_docx, pdf_path)
        print(f"✅ Generated: {name}_admission.pdf")
        
        # Optional: Delete the temporary docx to save space
        # os.remove(temp_docx)
        
        return pdf_path, True
    except Exception as e:
        print(f"❌ Failed for {name}: {e}")
        return temp_docx, False

# ---------- MAIN ----------
# Load student data
df = pd.read_excel("students.xlsx")
print(f"📊 Generating PDFs for {len(df)} students...\n")

success = 0
failed = 0
failed_names = []

for index, row in df.iterrows():
    pdf_path, generated = generate_pdf(row)
    if generated:
        success += 1
    else:
        failed += 1
        failed_names.append(row['Name'])
    time.sleep(0.5)

# Summary
print("\n" + "="*50)
print("📊 GENERATION SUMMARY")
print("="*50)
print(f"✅ Successful: {success}")
print(f"❌ Failed: {failed}")
print(f"📁 Output folder: {OUTPUT_DIR}")

if failed_names:
    print(f"\n❌ Failed students:")
    for name in failed_names:
        print(f"   - {name}")

print("\n✅ Done!")